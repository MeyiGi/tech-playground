import os
import math
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from sklearn.model_selection import train_test_split
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.impute import SimpleImputer
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.metrics import (
    accuracy_score,
    precision_score,
    recall_score,
    f1_score,
    roc_auc_score,
    roc_curve,
    confusion_matrix,
)
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from sklearn.neighbors import KNeighborsClassifier


# =========================
# CONFIG
# =========================
CSV_PATH = "titanic.csv"
OUTPUT_DOCX = "titanic_big_report.docx"
FIG_DIR = "report_figures"

os.makedirs(FIG_DIR, exist_ok=True)


# =========================
# DOCX HELPERS
# =========================
def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, font_size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_doc_margins(section, top=2, bottom=2, left=2.2, right=2.2):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def apply_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        s = doc.styles[style_name]
        s.font.name = "Times New Roman"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(18)

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.name = "Times New Roman"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r2.font.size = Pt(12)

    doc.add_paragraph("")


def add_paragraph(doc, text, bold=False, align="left"):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_image(doc, image_path, caption=None, width_inches=6.3):
    if not os.path.exists(image_path):
        add_paragraph(doc, f"Файл изображения не найден: {image_path}")
        return

    doc.add_picture(image_path, width=Inches(width_inches))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.italic = True
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(10)


def add_metrics_table(doc, metrics_df):
    headers = ["Модель", "Accuracy", "Precision", "Recall", "F1-score", "ROC-AUC"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(header_cells[i], h, bold=True, font_size=10)
        set_cell_shading(header_cells[i], "D9EAF7")

    for _, row in metrics_df.iterrows():
        cells = table.add_row().cells
        values = [
            row["Model"],
            f'{row["Accuracy"]:.4f}',
            f'{row["Precision"]:.4f}',
            f'{row["Recall"]:.4f}',
            f'{row["F1"]:.4f}',
            f'{row["ROC_AUC"]:.4f}',
        ]
        for i, v in enumerate(values):
            set_cell_text(cells[i], v, font_size=10)


# =========================
# PLOTTING HELPERS
# =========================
def save_figure(path):
    plt.tight_layout()
    plt.savefig(path, dpi=220, bbox_inches="tight")
    plt.close()


def plot_target_distribution(df):
    counts = df["Survived"].value_counts().sort_index()

    plt.figure(figsize=(8, 5))
    plt.bar(["No", "Yes"], counts.values)
    plt.title("Survival Distribution")
    plt.xlabel("Survived")
    plt.ylabel("Count")

    for i, v in enumerate(counts.values):
        plt.text(i, v + 5, str(v), ha="center")

    save_figure(os.path.join(FIG_DIR, "01_survival_distribution.png"))


def plot_sex_survival(df):
    grp = pd.crosstab(df["Sex"], df["Survived"])

    plt.figure(figsize=(8, 5))
    x = np.arange(len(grp.index))
    width = 0.35

    plt.bar(x - width / 2, grp[0].values, width, label="No")
    plt.bar(x + width / 2, grp[1].values, width, label="Yes")

    plt.xticks(x, grp.index)
    plt.xlabel("Sex")
    plt.ylabel("Count")
    plt.title("Survival by Sex")
    plt.legend()

    save_figure(os.path.join(FIG_DIR, "02_survival_by_sex.png"))


def plot_pclass_survival(df):
    grp = pd.crosstab(df["Pclass"], df["Survived"])

    plt.figure(figsize=(8, 5))
    x = np.arange(len(grp.index))
    width = 0.35

    plt.bar(x - width / 2, grp[0].values, width, label="No")
    plt.bar(x + width / 2, grp[1].values, width, label="Yes")

    plt.xticks(x, grp.index)
    plt.xlabel("Passenger Class")
    plt.ylabel("Count")
    plt.title("Survival by Passenger Class")
    plt.legend()

    save_figure(os.path.join(FIG_DIR, "03_survival_by_class.png"))


def plot_age_distribution(df):
    plt.figure(figsize=(8, 5))
    ages = df["Age"].dropna()
    plt.hist(ages, bins=25)
    plt.title("Age Distribution")
    plt.xlabel("Age")
    plt.ylabel("Frequency")
    save_figure(os.path.join(FIG_DIR, "04_age_distribution.png"))


def plot_fare_distribution(df):
    plt.figure(figsize=(8, 5))
    fares = df["Fare"].dropna()
    plt.hist(fares, bins=30)
    plt.title("Fare Distribution")
    plt.xlabel("Fare")
    plt.ylabel("Frequency")
    save_figure(os.path.join(FIG_DIR, "05_fare_distribution.png"))


def plot_correlation_heatmap(df):
    numeric_df = df.select_dtypes(include=[np.number]).copy()
    corr = numeric_df.corr()

    plt.figure(figsize=(9, 7))
    im = plt.imshow(corr, aspect="auto")
    plt.colorbar(im)
    plt.xticks(range(len(corr.columns)), corr.columns, rotation=45, ha="right")
    plt.yticks(range(len(corr.columns)), corr.columns)
    plt.title("Correlation Matrix")

    for i in range(len(corr.index)):
        for j in range(len(corr.columns)):
            plt.text(j, i, f"{corr.iloc[i, j]:.2f}", ha="center", va="center", fontsize=8)

    save_figure(os.path.join(FIG_DIR, "06_correlation_matrix.png"))


def plot_metrics_comparison(metrics_df):
    metric_cols = ["Accuracy", "Precision", "Recall", "F1", "ROC_AUC"]

    for metric in metric_cols:
        plt.figure(figsize=(8, 5))
        plt.bar(metrics_df["Model"], metrics_df[metric])
        plt.ylim(0, 1)
        plt.title(f"{metric} Comparison")
        plt.xlabel("Model")
        plt.ylabel(metric)

        for i, v in enumerate(metrics_df[metric]):
            plt.text(i, min(v + 0.01, 0.98), f"{v:.3f}", ha="center")

        save_figure(os.path.join(FIG_DIR, f"comparison_{metric.lower()}.png"))


def plot_feature_importance_random_forest(model_pipeline, feature_names):
    rf = model_pipeline.named_steps["model"]
    importances = rf.feature_importances_

    imp_df = pd.DataFrame({
        "Feature": feature_names,
        "Importance": importances
    }).sort_values("Importance", ascending=False).head(15)

    plt.figure(figsize=(10, 6))
    plt.barh(imp_df["Feature"][::-1], imp_df["Importance"][::-1])
    plt.title("Random Forest Feature Importance (Top 15)")
    plt.xlabel("Importance")
    plt.ylabel("Feature")
    save_figure(os.path.join(FIG_DIR, "rf_feature_importance.png"))


def plot_confusion_matrix(cm, labels, title, filename):
    plt.figure(figsize=(6, 5))
    plt.imshow(cm, interpolation="nearest")
    plt.title(title)
    plt.colorbar()
    tick_marks = np.arange(len(labels))
    plt.xticks(tick_marks, labels)
    plt.yticks(tick_marks, labels)
    plt.xlabel("Predicted")
    plt.ylabel("Actual")

    threshold = cm.max() / 2 if cm.max() > 0 else 0
    for i in range(cm.shape[0]):
        for j in range(cm.shape[1]):
            plt.text(
                j,
                i,
                str(cm[i, j]),
                ha="center",
                va="center",
                color="white" if cm[i, j] > threshold else "black",
                fontsize=12,
            )

    save_figure(os.path.join(FIG_DIR, filename))


def plot_roc_curve_single(fpr, tpr, roc_auc, title, filename):
    plt.figure(figsize=(6, 5))
    plt.plot(fpr, tpr, label=f"AUC = {roc_auc:.4f}")
    plt.plot([0, 1], [0, 1], linestyle="--")
    plt.xlabel("False Positive Rate")
    plt.ylabel("True Positive Rate")
    plt.title(title)
    plt.legend(loc="lower right")
    save_figure(os.path.join(FIG_DIR, filename))


def plot_prediction_distribution(y_prob, title, filename):
    plt.figure(figsize=(7, 5))
    plt.hist(y_prob, bins=20)
    plt.xlabel("Predicted Probability")
    plt.ylabel("Count")
    plt.title(title)
    save_figure(os.path.join(FIG_DIR, filename))


# =========================
# LOAD DATA
# =========================
df = pd.read_csv(CSV_PATH)

# =========================
# BASIC INFO
# =========================
missing_info = df.isnull().sum()
missing_info = missing_info[missing_info > 0].sort_values(ascending=False)

# =========================
# EDA PLOTS
# =========================
plot_target_distribution(df)
plot_sex_survival(df)
plot_pclass_survival(df)
plot_age_distribution(df)
plot_fare_distribution(df)
plot_correlation_heatmap(df)

# =========================
# FEATURES
# =========================
target_col = "Survived"
drop_cols = ["PassengerId", "Name", "Ticket", "Cabin"]

X = df.drop(columns=[c for c in drop_cols if c in df.columns] + [target_col])
y = df[target_col]

numeric_features = X.select_dtypes(include=[np.number]).columns.tolist()
categorical_features = X.select_dtypes(exclude=[np.number]).columns.tolist()

numeric_transformer = Pipeline(steps=[
    ("imputer", SimpleImputer(strategy="median")),
    ("scaler", StandardScaler()),
])

categorical_transformer = Pipeline(steps=[
    ("imputer", SimpleImputer(strategy="most_frequent")),
    ("onehot", OneHotEncoder(handle_unknown="ignore")),
])

preprocessor = ColumnTransformer(
    transformers=[
        ("num", numeric_transformer, numeric_features),
        ("cat", categorical_transformer, categorical_features),
    ]
)

X_train, X_test, y_train, y_test = train_test_split(
    X, y,
    test_size=0.2,
    random_state=42,
    stratify=y
)

# =========================
# MODELS
# =========================
models = {
    "Logistic Regression": LogisticRegression(max_iter=2000),
    "Random Forest": RandomForestClassifier(
        n_estimators=300,
        max_depth=None,
        random_state=42
    ),
    "KNN": KNeighborsClassifier(n_neighbors=7),
}

results = []
trained_pipelines = {}

for model_name, model in models.items():
    pipeline = Pipeline(steps=[
        ("preprocessor", preprocessor),
        ("model", model),
    ])

    pipeline.fit(X_train, y_train)
    y_pred = pipeline.predict(X_test)
    y_prob = pipeline.predict_proba(X_test)[:, 1]

    acc = accuracy_score(y_test, y_pred)
    prec = precision_score(y_test, y_pred, zero_division=0)
    rec = recall_score(y_test, y_pred, zero_division=0)
    f1 = f1_score(y_test, y_pred, zero_division=0)
    roc_auc = roc_auc_score(y_test, y_prob)

    cm = confusion_matrix(y_test, y_pred)
    fpr, tpr, _ = roc_curve(y_test, y_prob)

    safe_name = model_name.lower().replace(" ", "_")

    plot_confusion_matrix(
        cm,
        labels=["No", "Yes"],
        title=f"{model_name} - Confusion Matrix",
        filename=f"{safe_name}_cm.png"
    )

    plot_roc_curve_single(
        fpr,
        tpr,
        roc_auc,
        title=f"{model_name} - ROC Curve",
        filename=f"{safe_name}_roc.png"
    )

    plot_prediction_distribution(
        y_prob,
        title=f"{model_name} - Predicted Probability Distribution",
        filename=f"{safe_name}_prob_dist.png"
    )

    results.append({
        "Model": model_name,
        "Accuracy": acc,
        "Precision": prec,
        "Recall": rec,
        "F1": f1,
        "ROC_AUC": roc_auc,
    })

    trained_pipelines[model_name] = pipeline

metrics_df = pd.DataFrame(results).sort_values("ROC_AUC", ascending=False).reset_index(drop=True)

# comparison plots
plot_metrics_comparison(metrics_df)

# random forest feature importance
rf_pipeline = trained_pipelines["Random Forest"]
ohe = rf_pipeline.named_steps["preprocessor"].named_transformers_["cat"].named_steps["onehot"]

cat_feature_names = ohe.get_feature_names_out(categorical_features).tolist()
all_feature_names = numeric_features + cat_feature_names
plot_feature_importance_random_forest(rf_pipeline, all_feature_names)

best_model_name = metrics_df.iloc[0]["Model"]
best_model_row = metrics_df.iloc[0].to_dict()


# =========================
# CREATE DOCX
# =========================
doc = Document()
apply_base_style(doc)
set_doc_margins(doc.sections[0], top=2, bottom=2, left=2.3, right=2.3)

add_title(
    doc,
    "ОТЧЕТ ПО АНАЛИЗУ ДАННЫХ TITANIC",
    "Расширенный отчет с визуализациями, сравнением моделей и отдельными графиками для каждой модели"
)

doc.add_heading("1. Введение", level=1)
add_paragraph(
    doc,
    "В данном отчете проведен анализ датасета Titanic, выполнена предобработка данных, "
    "обучены несколько моделей машинного обучения и построены визуализации для сравнения их качества.",
    align="justify"
)

doc.add_heading("2. Общая информация о данных", level=1)
add_paragraph(doc, f"Количество строк: {df.shape[0]}")
add_paragraph(doc, f"Количество столбцов: {df.shape[1]}")
add_paragraph(doc, f"Целевая переменная: {target_col}")

doc.add_paragraph("")
add_paragraph(doc, "Пропущенные значения:", bold=True)
if len(missing_info) == 0:
    add_paragraph(doc, "Пропущенных значений не обнаружено.")
else:
    for col, val in missing_info.items():
        add_paragraph(doc, f"{col}: {val}")

doc.add_heading("3. Используемые признаки", level=1)
add_bullets(doc, [f for f in X.columns.tolist()])

doc.add_heading("4. Предобработка данных", level=1)
add_bullets(doc, [
    "Удалены малоинформативные поля: PassengerId, Name, Ticket, Cabin.",
    "Для числовых признаков пропуски заполнены медианой.",
    "Для категориальных признаков пропуски заполнены наиболее частым значением.",
    "Числовые признаки стандартизированы.",
    "Категориальные признаки преобразованы через One-Hot Encoding.",
    "Данные разделены на train/test в соотношении 80/20 со стратификацией.",
])

doc.add_heading("5. Исследовательский анализ данных (EDA)", level=1)
add_paragraph(doc, "Ниже приведены основные визуализации, отражающие структуру и закономерности датасета.")

eda_images = [
    ("01_survival_distribution.png", "Рисунок 1. Распределение выживших и невыживших пассажиров"),
    ("02_survival_by_sex.png", "Рисунок 2. Выживаемость по полу"),
    ("03_survival_by_class.png", "Рисунок 3. Выживаемость по классу билета"),
    ("04_age_distribution.png", "Рисунок 4. Распределение возраста"),
    ("05_fare_distribution.png", "Рисунок 5. Распределение стоимости билета"),
    ("06_correlation_matrix.png", "Рисунок 6. Матрица корреляций"),
]

for file_name, caption in eda_images:
    add_image(doc, os.path.join(FIG_DIR, file_name), caption=caption, width_inches=6.2)
    doc.add_paragraph("")

doc.add_heading("6. Обученные модели", level=1)
add_bullets(doc, [
    "Logistic Regression",
    "Random Forest",
    "K-Nearest Neighbors (KNN)",
])

doc.add_heading("7. Сравнение моделей по метрикам", level=1)
add_paragraph(
    doc,
    "Для оценки моделей использованы метрики Accuracy, Precision, Recall, F1-score и ROC-AUC."
)
add_metrics_table(doc, metrics_df)
doc.add_paragraph("")

comparison_images = [
    ("comparison_accuracy.png", "Рисунок 7. Сравнение Accuracy"),
    ("comparison_precision.png", "Рисунок 8. Сравнение Precision"),
    ("comparison_recall.png", "Рисунок 9. Сравнение Recall"),
    ("comparison_f1.png", "Рисунок 10. Сравнение F1-score"),
    ("comparison_roc_auc.png", "Рисунок 11. Сравнение ROC-AUC"),
]

for file_name, caption in comparison_images:
    add_image(doc, os.path.join(FIG_DIR, file_name), caption=caption, width_inches=5.8)
    doc.add_paragraph("")

doc.add_heading("8. Детальный анализ каждой модели", level=1)

for model_name in ["Logistic Regression", "Random Forest", "KNN"]:
    row = metrics_df[metrics_df["Model"] == model_name].iloc[0]
    safe_name = model_name.lower().replace(" ", "_")

    doc.add_heading(model_name, level=2)
    add_paragraph(
        doc,
        (
            f"Accuracy = {row['Accuracy']:.4f}, "
            f"Precision = {row['Precision']:.4f}, "
            f"Recall = {row['Recall']:.4f}, "
            f"F1-score = {row['F1']:.4f}, "
            f"ROC-AUC = {row['ROC_AUC']:.4f}"
        )
    )

    add_image(
        doc,
        os.path.join(FIG_DIR, f"{safe_name}_cm.png"),
        caption=f"{model_name}: матрица ошибок",
        width_inches=5.2
    )

    add_image(
        doc,
        os.path.join(FIG_DIR, f"{safe_name}_roc.png"),
        caption=f"{model_name}: ROC-кривая",
        width_inches=5.2
    )

    add_image(
        doc,
        os.path.join(FIG_DIR, f"{safe_name}_prob_dist.png"),
        caption=f"{model_name}: распределение предсказанных вероятностей",
        width_inches=5.4
    )

    doc.add_paragraph("")

doc.add_heading("9. Важность признаков Random Forest", level=1)
add_paragraph(
    doc,
    "Для модели Random Forest дополнительно построен график важности признаков."
)
add_image(
    doc,
    os.path.join(FIG_DIR, "rf_feature_importance.png"),
    caption="Рисунок 12. Важность признаков Random Forest",
    width_inches=6.0
)

doc.add_heading("10. Итоговый вывод", level=1)
add_paragraph(
    doc,
    (
        f"Лучшей моделью по метрике ROC-AUC стала модель '{best_model_name}'. "
        f"Ее показатели: Accuracy = {best_model_row['Accuracy']:.4f}, "
        f"Precision = {best_model_row['Precision']:.4f}, "
        f"Recall = {best_model_row['Recall']:.4f}, "
        f"F1-score = {best_model_row['F1']:.4f}, "
        f"ROC-AUC = {best_model_row['ROC_AUC']:.4f}."
    ),
    align="justify"
)

doc.add_heading("11. Рекомендации", level=1)
add_bullets(doc, [
    "Подобрать гиперпараметры моделей через GridSearchCV или RandomizedSearchCV.",
    "Добавить новые признаки, например FamilySize или IsAlone.",
    "Проверить кросс-валидацию для более устойчивой оценки качества.",
    "Сравнить результат с XGBoost или Gradient Boosting.",
])

doc.save(OUTPUT_DOCX)
print(f"Готово: {OUTPUT_DOCX}")
print(f"Графики сохранены в папке: {FIG_DIR}")